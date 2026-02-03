import os
import csv
import docx

class FileHandler:
    @staticmethod
    def read_file(file_path):
        """
        Reads the content of the file based on its extension.
        Returns a list of strings (lines) or raises an error.
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")

        ext = os.path.splitext(file_path)[1].lower()

        if ext == '.txt':
            return FileHandler._read_txt(file_path)
        elif ext == '.docx':
            return FileHandler._read_docx(file_path)
        elif ext == '.csv':
            return FileHandler._read_csv(file_path)
        else:
            raise ValueError(f"Unsupported file format: {ext}")

    @staticmethod
    def _read_txt(file_path):
        with open(file_path, 'r', encoding='utf-8') as f:
            return [line.strip() for line in f.readlines() if line.strip()]

    @staticmethod
    def _read_docx(file_path):
        doc = docx.Document(file_path)
        lines = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                lines.append(paragraph.text.strip())
        # Also check for tables in docx
        for table in doc.tables:
            for row in table.rows:
                # Naive approach: join cells with a separator that Parser can detect
                # Or return list of lists? For now, let's return text lines joined by ' - '
                # if it looks like a vocab pair.
                cell_texts = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cell_texts:
                    lines.append(" - ".join(cell_texts))
        return lines

    @staticmethod
    def _read_csv(file_path):
        lines = []
        with open(file_path, 'r', encoding='utf-8') as f:
            reader = csv.reader(f)
            for row in reader:
                # Filter empty strings
                row = [item.strip() for item in row if item.strip()]
                if row:
                    lines.append(" - ".join(row)) # Normalizing to a text format for the parser
        return lines
